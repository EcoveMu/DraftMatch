import os
import io
import re
import tempfile
import fitz  # PyMuPDF
import pdfplumber
import pytesseract
from PIL import Image
import docx

def extract_text_from_word(file):
    """從Word文件中提取文本段落和表格。"""
    # 將上傳的Word文件保存到臨時文件
    temp_dir = tempfile.mkdtemp()
    temp_file_path = os.path.join(temp_dir, "temp.docx")
    with open(temp_file_path, "wb") as f:
        f.write(file.getvalue())
    # 打開Word文件
    doc = docx.Document(temp_file_path)
    paragraphs = []
    tables = []
    # 提取段落
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append({
                "index": len(paragraphs),
                "content": text,
                "type": "paragraph",
                "style": para.style.name if para.style else "Normal"
            })
    # 提取表格
    for i, table in enumerate(doc.tables):
        # 提取表格數據
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        # 跳過完全空白的表格
        if any(any(cell for cell in row) for row in table_data):
            # 嘗試獲取表格標題（表格前最近的段落中包含"表"或"Table"的作為標題）
            title = ""
            if paragraphs:
                for j in range(len(paragraphs)-1, -1, -1):
                    if "表" in paragraphs[j]["content"] or "Table" in paragraphs[j]["content"]:
                        title = paragraphs[j]["content"]
                        break
            tables.append({
                "index": len(tables),
                "title": title,
                "data": table_data,
                "type": "table"
            })
    # 刪除臨時文件
    os.remove(temp_file_path)
    os.rmdir(temp_dir)
    return {"paragraphs": paragraphs, "tables": tables}

def extract_text_from_pdf_with_pdfplumber(file):
    """使用pdfplumber從PDF文件中提取文本段落和表格。"""
    # 將上傳的PDF保存到臨時文件
    temp_dir = tempfile.mkdtemp()
    temp_file_path = os.path.join(temp_dir, "temp.pdf")
    with open(temp_file_path, "wb") as f:
        f.write(file.getvalue())
    pdf = pdfplumber.open(temp_file_path)
    all_paragraphs = []
    all_tables = []
    for i, page in enumerate(pdf.pages):
        # 提取頁面文本並分段
        text = page.extract_text()
        if text:
            page_paragraphs = text.split('\n\n')
            for para in page_paragraphs:
                para = para.strip()
                if para:
                    all_paragraphs.append({
                        "index": len(all_paragraphs),
                        "content": para,
                        "page": i + 1,
                        "type": "paragraph"
                    })
        # 提取頁面中的表格
        tables = page.extract_tables()
        for table in tables:
            # 過濾掉空表格
            if table and any(any(cell for cell in row) for row in table):
                title = ""
                if all_paragraphs:
                    for k in range(len(all_paragraphs)-1, -1, -1):
                        if "表" in all_paragraphs[k]["content"] or "Table" in all_paragraphs[k]["content"]:
                            title = all_paragraphs[k]["content"]
                            break
                all_tables.append({
                    "index": len(all_tables),
                    "title": title,
                    "data": table,
                    "page": i + 1,
                    "type": "table"
                })
    pdf.close()
    os.remove(temp_file_path)
    os.rmdir(temp_dir)
    return {"paragraphs": all_paragraphs, "tables": all_tables}

def extract_text_from_pdf_with_pymupdf(file):
    """使用PyMuPDF從PDF文件中提取文本段落和表格。"""
    temp_dir = tempfile.mkdtemp()
    temp_file_path = os.path.join(temp_dir, "temp.pdf")
    with open(temp_file_path, "wb") as f:
        f.write(file.getvalue())
    doc = fitz.open(temp_file_path)
    all_paragraphs = []
    all_tables = []
    for i in range(len(doc)):
        page = doc[i]
        # 提取文本並分段
        text = page.get_text()
        if text:
            page_paragraphs = text.split('\n\n')
            for para in page_paragraphs:
                para = para.strip()
                if para:
                    all_paragraphs.append({
                        "index": len(all_paragraphs),
                        "content": para,
                        "page": i + 1,
                        "type": "paragraph"
                    })
        # 嘗試使用簡單啟發式方法提取表格
        blocks = page.get_text("blocks")
        for block in blocks:
            if block[4].count("\n") > 2 and "\t" in block[4]:
                table_text = block[4]
                rows = table_text.split("\n")
                table_data = []
                for row in rows:
                    if row.strip():
                        cells = row.split("\t")
                        table_data.append([cell.strip() for cell in cells])
                if table_data and len(table_data) > 1 and len(table_data[0]) > 1:
                    title = ""
                    if all_paragraphs:
                        for k in range(len(all_paragraphs)-1, -1, -1):
                            if "表" in all_paragraphs[k]["content"] or "Table" in all_paragraphs[k]["content"]:
                                title = all_paragraphs[k]["content"]
                                break
                    all_tables.append({
                        "index": len(all_tables),
                        "title": title,
                        "data": table_data,
                        "page": i + 1,
                        "type": "table"
                    })
    doc.close()
    os.remove(temp_file_path)
    os.rmdir(temp_dir)
    return {"paragraphs": all_paragraphs, "tables": all_tables}

def extract_text_from_pdf_with_page_info(pdf_file):
    """從PDF文件中提取文本段落，保留每個段落的頁碼信息。"""
    temp_dir = tempfile.mkdtemp()
    temp_pdf_path = os.path.join(temp_dir, "temp.pdf")
    with open(temp_pdf_path, "wb") as f:
        f.write(pdf_file.getvalue())
    doc = fitz.open(temp_pdf_path)
    paragraphs = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        if text:
            page_paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            for para_text in page_paragraphs:
                paragraphs.append({
                    "index": len(paragraphs),
                    "content": para_text,
                    "type": "paragraph",
                    "page": page_num + 1
                })
    doc.close()
    os.remove(temp_pdf_path)
    os.rmdir(temp_dir)
    return paragraphs

def extract_text_from_pdf_with_ocr(file, ocr_engine="tesseract", ocr_instance=None):
    """使用OCR從PDF文件中提取文本段落和表格。"""
    temp_dir = tempfile.mkdtemp()
    temp_file_path = os.path.join(temp_dir, "temp.pdf")
    with open(temp_file_path, "wb") as f:
        f.write(file.getvalue())
    doc = fitz.open(temp_file_path)
    all_paragraphs = []
    all_tables = []
    print("正在使用OCR提取文本，這可能需要一些時間...")
    for i in range(len(doc)):
        print(f"正在OCR處理頁面 {i+1}...")
        page = doc[i]
        pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
        img_path = os.path.join(temp_dir, f"page_{i+1}.png")
        pix.save(img_path)
        
        # 根據指定的OCR引擎提取文本
        text = None
        if ocr_instance and ocr_instance.is_available():
            if ocr_engine.lower() in ["qwen_builtin", "qwen"]:
                text, _ = ocr_instance.extract_text_from_pdf_page(temp_file_path, i)
                if isinstance(text, str) and text.startswith("從PDF提取文本時出錯"):
                    text = ocr_instance.extract_text_from_image(img_path)
            elif ocr_engine.lower() in ["easyocr", "tesseract", "ocr_custom"]:
                text = ocr_instance.extract_text_from_image(img_path)
        
        # 如果沒有OCR實例或OCR提取失敗，使用默認的Tesseract
        if text is None or not isinstance(text, str) or len(text.strip()) == 0:
            try:
                # 嘗試使用Tesseract作為備選
                img = Image.open(img_path)
                text = pytesseract.image_to_string(img, lang='chi_tra+eng')
                print(f"使用備選Tesseract OCR提取文本")
            except Exception as e:
                print(f"備選OCR失敗: {str(e)}")
                text = ""
        
        print(f"OCR頁面 {i+1} 提取的文本長度: {len(text) if text else 0}")
        
        # 保存提取到的段落
        if text:
            page_paragraphs = text.split('\n\n')
            for para in page_paragraphs:
                para = para.strip()
                if para:
                    all_paragraphs.append({
                        "index": len(all_paragraphs),
                        "content": para,
                        "page": i + 1,
                        "type": "paragraph"
                    })
        
        # 提取表格
        if ocr_instance and ocr_instance.is_available():
            tables = ocr_instance.extract_tables_from_image(img_path)
            if isinstance(tables, list):
                for table in tables:
                    if table and len(table) > 1 and len(table[0]) > 1:
                        title = ""
                        if all_paragraphs:
                            for k in range(len(all_paragraphs)-1, -1, -1):
                                if "表" in all_paragraphs[k]["content"] or "Table" in all_paragraphs[k]["content"]:
                                    title = all_paragraphs[k]["content"]
                                    break
                        all_tables.append({
                            "index": len(all_tables),
                            "title": title,
                            "data": table,
                            "page": i + 1,
                            "type": "table"
                        })
        else:
            # 使用簡單的啟發式從OCR文本中檢測表格
            lines = text.split('\n') if text else []
            table_start = -1
            table_end = -1
            for j, line in enumerate(lines):
                if line.strip() and (line.count('\t') > 1 or line.count('  ') > 2):
                    if table_start == -1:
                        table_start = j
                    table_end = j
                elif table_start != -1 and j - table_end > 1:
                    table_lines = lines[table_start:table_end+1]
                    table_data = []
                    for table_line in table_lines:
                        if table_line.strip():
                            cells = re.split(r'\t|  +', table_line)
                            table_data.append([cell.strip() for cell in cells if cell.strip()])
                    if table_data and len(table_data) > 1 and len(table_data[0]) > 1:
                        title = ""
                        if table_start > 0:
                            title = lines[table_start-1].strip()
                        all_tables.append({
                            "index": len(all_tables),
                            "title": title,
                            "data": table_data,
                            "page": i + 1,
                            "type": "table"
                        })
                    table_start = -1
                    table_end = -1
    
    # 清理臨時文件
    doc.close()
    for i in range(len(doc)):
        img_path = os.path.join(temp_dir, f"page_{i+1}.png")
        if os.path.exists(img_path):
            os.remove(img_path)
    os.remove(temp_file_path)
    os.rmdir(temp_dir)
    
    return {"paragraphs": all_paragraphs, "tables": all_tables}

def enhanced_extract_tables_from_pdf(file) -> list:
    """Enhanced function to extract tables from PDF with better table detection and title matching"""
    # Create temp file
    temp_dir = tempfile.mkdtemp()
    temp_file_path = os.path.join(temp_dir, "temp.pdf")
    with open(temp_file_path, "wb") as f:
        f.write(file.getvalue())

    pdf = pdfplumber.open(temp_file_path)
    tables = []
    paragraphs = []

    # First pass: extract all paragraphs for title matching
    for page_num, page in enumerate(pdf.pages, start=1):
        text = page.extract_text()
        if text:
            page_paragraphs = text.split('\n\n')
            for para in page_paragraphs:
                para = para.strip()
                if para:
                    paragraphs.append({
                        "content": para,
                        "page": page_num
                    })

    # Second pass: extract tables with better title matching
    for page_num, page in enumerate(pdf.pages, start=1):
        page_tables = page.extract_tables()
        for table in page_tables:
            # Filter empty tables
            if table and any(any(cell and cell.strip() for cell in row) for row in table):
                # Look for table title in nearby paragraphs
                title = ""
                current_page_paras = [p for p in paragraphs if p["page"] <= page_num]
                
                # Check last few paragraphs for title
                for para in reversed(current_page_paras[-5:]):
                    content = para["content"]
                    if ("表" in content or "Table" in content) and len(content) < 100:
                        title = content
                        break

                # Clean table data
                cleaned_table = [
                    [cell.strip() if cell else "" for cell in row]
                    for row in table
                ]

                tables.append({
                    "page": page_num,
                    "index": len(tables),
                    "title": title,
                    "data": cleaned_table,
                    "type": "table"
                })

    pdf.close()
    # Cleanup
    os.remove(temp_file_path)
    os.rmdir(temp_dir)
    return tables

def extract_and_process_documents(word_file, pdf_file, use_ocr=True, ocr_engine="tesseract", ocr_instance=None):
    """提取並處理Word和PDF文件內容。"""
    # 提取Word文件內容
    word_data = extract_text_from_word(word_file)
    # 清理Word段落文本
    for para in word_data["paragraphs"]:
        para["content"] = clean_extracted_text(para["content"])
    # 提取PDF文件內容（使用PyMuPDF和pdfplumber）
    pdf_data_pymupdf = extract_text_from_pdf_with_pymupdf(pdf_file)
    pdf_data_pdfplumber = extract_text_from_pdf_with_pdfplumber(pdf_file)
    # 打印提取的文本長度資訊（調試用途）
    try:
        total_pages = len(fitz.open(stream=pdf_file.getvalue(), filetype="pdf"))
        print(f"PDF總頁數: {total_pages}")
        for para in pdf_data_pymupdf["paragraphs"]:
            print(f"PyMuPDF頁面 {para['page']} 提取的文本長度: {len(para['content'])}")
        for para in pdf_data_pdfplumber["paragraphs"]:
            print(f"pdfplumber頁面 {para['page']} 提取的文本長度: {len(para['content'])}")
    except Exception:
        pass
    # 如果使用OCR，使用OCR引擎提取補充的PDF文本
    pdf_data_ocr = {"paragraphs": [], "tables": []}
    if use_ocr:
        pdf_data_ocr = extract_text_from_pdf_with_ocr(pdf_file, ocr_engine, ocr_instance)
    # Extract PDF content with enhanced table extraction
    pdf_tables = enhanced_extract_tables_from_pdf(pdf_file)
    # 合併提取的段落和表格
    combined_paragraphs = pdf_data_pymupdf["paragraphs"] + pdf_data_pdfplumber["paragraphs"] + pdf_data_ocr.get("paragraphs", [])
    combined_tables = pdf_tables + pdf_data_pymupdf["tables"] + pdf_data_ocr.get("tables", [])
    # 去除重複的段落
    unique_paragraphs = []
    seen_contents = set()
    for para in combined_paragraphs:
        content = para["content"]
        if content not in seen_contents and len(content) > 5:  # 忽略極短的段落
            seen_contents.add(content)
            unique_paragraphs.append(para)
    # 去除重複的表格
    unique_tables = []
    seen_tables = set()
    for table in combined_tables:
        table_str = str(table["data"])
        if table_str not in seen_tables:
            seen_tables.add(table_str)
            unique_tables.append(table)
    # 清理PDF段落文本
    for para in unique_paragraphs:
        para["content"] = clean_extracted_text(para["content"])
    print(f"提取的總段落數: {len(unique_paragraphs)}")
    return word_data, {"paragraphs": unique_paragraphs, "tables": unique_tables}

def normalize_chinese_text(text):
    """標準化中文文本（全形轉半形，移除多餘空白）。"""
    # 全形轉半形
    import unicodedata
    text = unicodedata.normalize('NFKC', text)
    # 移除多餘空白
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def clean_extracted_text(text):
    """清理提取的文本，移除控制字符並標準化換行等。"""
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    # 多個連續換行替換為兩個換行
    text = re.sub(r'\n{3,}', '\n\n', text)
    # 標準化中文文本
    text = normalize_chinese_text(text)
    return text

def extract_page_numbers(text):
    """從給定文本中提取頁碼。"""
    page_patterns = [
        r'第\s*(\d+)\s*頁',
        r'Page\s*(\d+)',
        r'P\.?\s*(\d+)',
        r'頁\s*(\d+)',
        r'-\s*(\d+)\s*-'
    ]
    for pattern in page_patterns:
        match = re.search(pattern, text)
        if match:
            try:
                return int(match.group(1))
            except:
                continue
    return None

def preprocess_text(text, ignore_options):
    """根據忽略選項對文本進行預處理。"""
    if ignore_options.get("ignore_space", False):
        text = re.sub(r'\s+', '', text)
    if ignore_options.get("ignore_punctuation", False):
        text = re.sub(r'[^\w\s]', '', text)
    if ignore_options.get("ignore_case", False):
        text = text.lower()
    if ignore_options.get("ignore_newline", False):
        text = text.replace('\n', ' ')
        text = re.sub(r'\s+', ' ', text)
    return text

def detect_encoding(text):
    """檢測文本內容的語言或編碼類型。"""
    if re.search(r'[\uFFFD\u25A1\u2753]', text):
        return "亂碼"
    if re.search(r'[\u4e00-\u9fff]', text):
        return "中文"
    if re.search(r'[\u3040-\u309F\u30A0-\u30FF]', text):
        return "日文"
    if re.search(r'[\uAC00-\uD7AF]', text):
        return "韓文"
    return "英文"
