import os
import tempfile
import docx
import fitz  # PyMuPDF
import pdfplumber
import pytesseract
from PIL import Image
import pandas as pd
import numpy as np
import re
import io
from qwen_ocr import QwenOCR

def extract_text_from_word(file):
    """從Word文件中提取文本"""
    # 創建臨時文件
    temp_dir = tempfile.mkdtemp()
    temp_file_path = os.path.join(temp_dir, "temp.docx")
    
    # 保存上傳的文件
    with open(temp_file_path, "wb") as f:
        f.write(file.getvalue())
    
    # 打開Word文件
    doc = docx.Document(temp_file_path)
    
    # 提取段落
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append({
                "index": i,
                "content": text,
                "style": para.style.name if para.style else "Normal"
            })
    
    # 提取表格
    tables = []
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        
        # 嘗試提取表格標題（通常是表格前的段落）
        title = ""
        if i > 0 and len(paragraphs) > 0:
            # 尋找表格前的段落作為標題
            for j in range(len(paragraphs)-1, -1, -1):
                if "表" in paragraphs[j]["content"] or "Table" in paragraphs[j]["content"]:
                    title = paragraphs[j]["content"]
                    break
        
        tables.append({
            "index": i,
            "title": title,
            "data": table_data
        })
    
    # 刪除臨時文件
    os.remove(temp_file_path)
    os.rmdir(temp_dir)
    
    return {
        "paragraphs": paragraphs,
        "tables": tables
    }

def extract_text_from_pdf_with_pdfplumber(file):
    """使用pdfplumber從PDF文件中提取文本"""
    # 創建臨時文件
    temp_dir = tempfile.mkdtemp()
    temp_file_path = os.path.join(temp_dir, "temp.pdf")
    
    # 保存上傳的文件
    with open(temp_file_path, "wb") as f:
        f.write(file.getvalue())
    
    # 打開PDF文件
    pdf = pdfplumber.open(temp_file_path)
    
    # 提取段落和表格
    all_paragraphs = []
    all_tables = []
    
    for i, page in enumerate(pdf.pages):
        # 提取文本
        text = page.extract_text()
        if text:
            # 分割成段落
            page_paragraphs = text.split('\n\n')
            for j, para in enumerate(page_paragraphs):
                para = para.strip()
                if para:
                    all_paragraphs.append({
                        "index": len(all_paragraphs),
                        "content": para,
                        "page": i + 1
                    })
        
        # 提取表格
        tables = page.extract_tables()
        for j, table in enumerate(tables):
            # 過濾空表格
            if table and any(any(cell for cell in row) for row in table):
                # 嘗試提取表格標題
                title = ""
                if len(all_paragraphs) > 0:
                    # 尋找表格前的段落作為標題
                    for k in range(len(all_paragraphs)-1, -1, -1):
                        if "表" in all_paragraphs[k]["content"] or "Table" in all_paragraphs[k]["content"]:
                            title = all_paragraphs[k]["content"]
                            break
                
                all_tables.append({
                    "index": len(all_tables),
                    "title": title,
                    "data": table,
                    "page": i + 1
                })
    
    pdf.close()
    
    # 刪除臨時文件
    os.remove(temp_file_path)
    os.rmdir(temp_dir)
    
    return {
        "paragraphs": all_paragraphs,
        "tables": all_tables
    }

def extract_text_from_pdf_with_pymupdf(file):
    """使用PyMuPDF從PDF文件中提取文本"""
    # 創建臨時文件
    temp_dir = tempfile.mkdtemp()
    temp_file_path = os.path.join(temp_dir, "temp.pdf")
    
    # 保存上傳的文件
    with open(temp_file_path, "wb") as f:
        f.write(file.getvalue())
    
    # 打開PDF文件
    doc = fitz.open(temp_file_path)
    
    # 提取段落和表格
    all_paragraphs = []
    all_tables = []
    
    for i in range(len(doc)):
        page = doc[i]
        
        # 提取文本
        text = page.get_text()
        if text:
            # 分割成段落
            page_paragraphs = text.split('\n\n')
            for j, para in enumerate(page_paragraphs):
                para = para.strip()
                if para:
                    all_paragraphs.append({
                        "index": len(all_paragraphs),
                        "content": para,
                        "page": i + 1
                    })
        
        # 提取表格（PyMuPDF的表格提取功能有限，這裡使用簡單的啟發式方法）
        # 尋找可能的表格區域
        blocks = page.get_text("blocks")
        for block in blocks:
            if block[4].count("\n") > 2 and block[4].count("\t") > 0:
                # 可能是表格
                table_text = block[4]
                rows = table_text.split("\n")
                table_data = []
                for row in rows:
                    if row.strip():
                        cells = row.split("\t")
                        table_data.append([cell.strip() for cell in cells])
                
                if table_data and len(table_data) > 1 and len(table_data[0]) > 1:
                    # 嘗試提取表格標題
                    title = ""
                    if len(all_paragraphs) > 0:
                        # 尋找表格前的段落作為標題
                        for k in range(len(all_paragraphs)-1, -1, -1):
                            if "表" in all_paragraphs[k]["content"] or "Table" in all_paragraphs[k]["content"]:
                                title = all_paragraphs[k]["content"]
                                break
                    
                    all_tables.append({
                        "index": len(all_tables),
                        "title": title,
                        "data": table_data,
                        "page": i + 1
                    })
    
    doc.close()
    
    # 刪除臨時文件
    os.remove(temp_file_path)
    os.rmdir(temp_dir)
    
    return {
        "paragraphs": all_paragraphs,
        "tables": all_tables
    }

def extract_text_from_pdf_with_ocr(file, ocr_engine="tesseract", ocr_instance=None):
    """使用OCR從PDF文件中提取文本"""
    # 創建臨時文件
    temp_dir = tempfile.mkdtemp()
    temp_file_path = os.path.join(temp_dir, "temp.pdf")
    
    # 保存上傳的文件
    with open(temp_file_path, "wb") as f:
        f.write(file.getvalue())
    
    # 打開PDF文件
    doc = fitz.open(temp_file_path)
    
    # 提取段落和表格
    all_paragraphs = []
    all_tables = []
    
    print("正在使用OCR提取文本，這可能需要一些時間...")
    
    for i in range(len(doc)):
        print(f"正在OCR處理頁面 {i+1}...")
        page = doc[i]
        
        # 將頁面轉換為圖像
        pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
        img_path = os.path.join(temp_dir, f"page_{i+1}.png")
        pix.save(img_path)
        
        # 使用OCR提取文本
        if ocr_engine == "tesseract":
            # 使用Tesseract OCR
            img = Image.open(img_path)
            text = pytesseract.image_to_string(img, lang='chi_tra+eng')
        elif ocr_engine == "Qwen" and ocr_instance:
            # 使用Qwen OCR
            text, _ = ocr_instance.extract_text_from_pdf_page(temp_file_path, i)
            if isinstance(text, str) and text.startswith("從PDF提取文本時出錯"):
                # 如果提取失敗，嘗試直接從圖像提取
                text = ocr_instance.extract_text_from_image(img_path)
        else:
            # 默認使用Tesseract
            img = Image.open(img_path)
            text = pytesseract.image_to_string(img, lang='chi_tra+eng')
        
        print(f"OCR頁面 {i+1} 提取的文本長度: {len(text)}")
        
        if text:
            # 分割成段落
            page_paragraphs = text.split('\n\n')
            for j, para in enumerate(page_paragraphs):
                para = para.strip()
                if para:
                    all_paragraphs.append({
                        "index": len(all_paragraphs),
                        "content": para,
                        "page": i + 1
                    })
        
        # 提取表格（使用OCR提取表格比較困難，這裡使用簡單的啟發式方法）
        if ocr_engine == "Qwen" and ocr_instance:
            # 使用Qwen提取表格
            tables = ocr_instance.extract_tables_from_image(img_path)
            if isinstance(tables, list):
                for j, table in enumerate(tables):
                    if table and len(table) > 1 and len(table[0]) > 1:
                        # 嘗試提取表格標題
                        title = ""
                        if len(all_paragraphs) > 0:
                            # 尋找表格前的段落作為標題
                            for k in range(len(all_paragraphs)-1, -1, -1):
                                if "表" in all_paragraphs[k]["content"] or "Table" in all_paragraphs[k]["content"]:
                                    title = all_paragraphs[k]["content"]
                                    break
                        
                        all_tables.append({
                            "index": len(all_tables),
                            "title": title,
                            "data": table,
                            "page": i + 1
                        })
        else:
            # 使用啟發式方法尋找表格
            lines = text.split('\n')
            table_start = -1
            table_end = -1
            for j, line in enumerate(lines):
                # 如果一行包含多個空格或製表符，可能是表格的一部分
                if line.strip() and (line.count('\t') > 1 or line.count('  ') > 2):
                    if table_start == -1:
                        table_start = j
                    table_end = j
                elif table_start != -1 and j - table_end > 1:
                    # 表格結束
                    table_lines = lines[table_start:table_end+1]
                    table_data = []
                    for table_line in table_lines:
                        if table_line.strip():
                            cells = re.split(r'\t|  +', table_line)
                            table_data.append([cell.strip() for cell in cells if cell.strip()])
                    
                    if table_data and len(table_data) > 1 and len(table_data[0]) > 1:
                        # 嘗試提取表格標題
                        title = ""
                        if table_start > 0:
                            title = lines[table_start-1].strip()
                        
                        all_tables.append({
                            "index": len(all_tables),
                            "title": title,
                            "data": table_data,
                            "page": i + 1
                        })
                    
                    table_start = -1
                    table_end = -1
    
    doc.close()
    
    # 刪除臨時文件
    for i in range(len(doc)):
        img_path = os.path.join(temp_dir, f"page_{i+1}.png")
        if os.path.exists(img_path):
            os.remove(img_path)
    
    os.remove(temp_file_path)
    os.rmdir(temp_dir)
    
    return {
        "paragraphs": all_paragraphs,
        "tables": all_tables
    }

def extract_and_process_documents(word_file, pdf_file, use_ocr=True, ocr_engine="tesseract", ocr_instance=None):
    """提取和處理文件內容"""
    # 提取Word文件內容
    word_data = extract_text_from_word(word_file)
    
    # 提取PDF文件內容
    pdf_data_pymupdf = extract_text_from_pdf_with_pymupdf(pdf_file)
    pdf_data_pdfplumber = extract_text_from_pdf_with_pdfplumber(pdf_file)
    
    # 打印提取的文本長度
    print(f"PDF總頁數: {len(fitz.open(io.BytesIO(pdf_file.getvalue())))}")
    for i, para in enumerate(pdf_data_pymupdf["paragraphs"]):
        if i < len(fitz.open(io.BytesIO(pdf_file.getvalue()))):
            print(f"頁面 {para['page']} 提取的文本長度: {len(para['content'])}")
    
    for i, para in enumerate(pdf_data_pdfplumber["paragraphs"]):
        if i < len(pdfplumber.open(io.BytesIO(pdf_file.getvalue())).pages):
            print(f"pdfplumber頁面 {para['page']} 提取的文本長度: {len(para['content'])}")
    
    # 如果使用OCR，則使用OCR提取文本
    if use_ocr:
        pdf_data_ocr = extract_text_from_pdf_with_ocr(pdf_file, ocr_engine, ocr_instance)
        
        # 合併提取的文本
        pdf_data = {
            "paragraphs": pdf_data_pymupdf["paragraphs"] + pdf_data_pdfplumber["paragraphs"] + pdf_data_ocr["paragraphs"],
            "tables": pdf_data_pymupdf["tables"] + pdf_data_pdfplumber["tables"] + pdf_data_ocr["tables"]
        }
    else:
        # 合併提取的文本
        pdf_data = {
            "paragraphs": pdf_data_pymupdf["paragraphs"] + pdf_data_pdfplumber["paragraphs"],
            "tables": pdf_data_pymupdf["tables"] + pdf_data_pdfplumber["tables"]
        }
    
    # 去除重複的段落
    unique_paragraphs = []
    seen_contents = set()
    
    for para in pdf_data["paragraphs"]:
        content = para["content"]
        if content not in seen_contents and len(content) > 5:  # 忽略太短的段落
            seen_contents.add(content)
            unique_paragraphs.append(para)
    
    pdf_data["paragraphs"] = unique_paragraphs
    
    # 去除重複的表格
    unique_tables = []
    seen_tables = set()
    
    for table in pdf_data["tables"]:
        # 將表格數據轉換為字符串進行比較
        table_str = str(table["data"])
        if table_str not in seen_tables:
            seen_tables.add(table_str)
            unique_tables.append(table)
    
    pdf_data["tables"] = unique_tables
    
    print(f"提取的總段落數: {len(pdf_data['paragraphs'])}")
    
    return word_data, pdf_data

# 測試函數
def test_extraction(word_file_path, pdf_file_path):
    """測試文本提取功能"""
    with open(word_file_path, "rb") as word_file, open(pdf_file_path, "rb") as pdf_file:
        word_data, pdf_data = extract_and_process_documents(word_file, pdf_file)
        
        print(f"Word文件段落數: {len(word_data['paragraphs'])}")
        print(f"Word文件表格數: {len(word_data['tables'])}")
        print(f"PDF文件段落數: {len(pdf_data['paragraphs'])}")
        print(f"PDF文件表格數: {len(pdf_data['tables'])}")
        
        # 打印一些示例段落
        print("\nWord文件段落示例:")
        for i, para in enumerate(word_data["paragraphs"][:3]):
            print(f"{i+1}. {para['content'][:100]}...")
        
        print("\nPDF文件段落示例:")
        for i, para in enumerate(pdf_data["paragraphs"][:3]):
            print(f"{i+1}. {para['content'][:100]}...")
        
        return word_data, pdf_data

# 主函數
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 3:
        print("用法: python text_extraction.py <word_file_path> <pdf_file_path>")
        sys.exit(1)
    
    word_file_path = sys.argv[1]
    pdf_file_path = sys.argv[2]
    test_extraction(word_file_path, pdf_file_path)
