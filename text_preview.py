import streamlit as st
import pandas as pd
from docx import Document
import fitz  # PyMuPDF
from qwen_ocr import QwenOCR
import os
import io
from PIL import Image
import numpy as np
import tempfile

class TextPreview:
    def __init__(self):
        self.qwen_ocr = QwenOCR()
        
    def extract_word_content(self, word_file):
        """從 Word 文件中提取內容，如果沒有文字則使用 OCR，提取所有元素包括標題、目錄等"""
        doc = Document(word_file)
        paragraphs = []
        paragraph_index = 0
        
        # 提取文檔屬性和元數據
        properties = doc.core_properties
        if properties.title:
            paragraphs.append({
                'index': paragraph_index,
                'content': f"標題: {properties.title}",
                'type': 'metadata'
            })
            paragraph_index += 1
        
        # 首先收集所有部分和節
        for i, section in enumerate(doc.sections):
            # 提取頁眉
            if section.header:
                for p in section.header.paragraphs:
                    if p.text.strip():
                        paragraphs.append({
                            'index': paragraph_index,
                            'content': f"頁眉: {p.text.strip()}",
                            'type': 'header'
                        })
                        paragraph_index += 1
            
            # 提取頁腳
            if section.footer:
                for p in section.footer.paragraphs:
                    if p.text.strip():
                        paragraphs.append({
                            'index': paragraph_index,
                            'content': f"頁腳: {p.text.strip()}",
                            'type': 'footer'
                        })
                        paragraph_index += 1
        
        # 處理所有段落，標記其樣式和層級
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
                
            # 分析段落樣式
            style_name = para.style.name if para.style else "Normal"
            paragraph_type = 'paragraph'
            
            # 識別標題
            if style_name.startswith('Heading') or '標題' in style_name:
                paragraph_type = 'heading'
            # 識別目錄
            elif 'TOC' in style_name or '目錄' in style_name:
                paragraph_type = 'toc'
            # 識別清單項目
            elif para.style.name.startswith('List'):
                paragraph_type = 'list'
            
            # 添加到結果中
            paragraphs.append({
                'index': paragraph_index,
                'content': text,
                'type': paragraph_type,
                'style': style_name
            })
            paragraph_index += 1
        
        # 提取表格內容 (作為文本)
        for i, table in enumerate(doc.tables):
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(" | ".join(row_text))
            
            if table_text:
                paragraphs.append({
                    'index': paragraph_index,
                    'content': "\n".join(table_text),
                    'type': 'table_text'
                })
                paragraph_index += 1
        
        # 如果沒有提取到任何內容，使用 OCR
        if not paragraphs:
            st.info("Word 文件中沒有找到足夠的文字內容，將使用 OCR 進行提取...")
            
            # 將 Word 轉換為 PDF
            pdf_bytes = io.BytesIO()
            doc.save(pdf_bytes)
            pdf_bytes.seek(0)
            
            # 使用 PyMuPDF 讀取 PDF
            pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                # 將頁面轉換為圖片
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # 將圖片保存為臨時文件
                temp_dir = tempfile.mkdtemp()
                temp_img_path = os.path.join(temp_dir, f"page_{page_num}.png")
                img.save(temp_img_path)
                
                # 使用 QwenOCR 提取文字
                text = self.qwen_ocr.extract_text_from_image(temp_img_path)
                
                # 清理臨時文件
                os.remove(temp_img_path)
                os.rmdir(temp_dir)
                
                if text and not text.startswith("提取文本時出錯"):
                    paragraphs.append({
                        'index': paragraph_index,
                        'content': text.strip(),
                        'type': 'ocr_text'
                    })
                    paragraph_index += 1
            
            pdf_doc.close()
        
        return paragraphs
    
    def extract_pdf_content(self, pdf_file):
        """從 PDF 文件中提取內容，如果沒有文字則使用 OCR"""
        doc = fitz.open(pdf_file)
        paragraphs = []
        
        # 先嘗試直接提取文字
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text().strip()
            
            if text:
                paragraphs.append({
                    'index': page_num,
                    'content': text,
                    'page': page_num + 1
                })
            else:
                # 如果沒有文字，使用 OCR
                st.info(f"PDF 第 {page_num + 1} 頁沒有找到文字內容，將使用 OCR 進行提取...")
                
                # 將頁面轉換為圖片
                pix = page.get_pixmap()
                
                # 將圖像保存為臨時文件
                temp_dir = tempfile.mkdtemp()
                temp_img_path = os.path.join(temp_dir, f"page_{page_num}.png")
                pix.save(temp_img_path)
                
                # 使用 QwenOCR 提取文字
                text = self.qwen_ocr.extract_text_from_image(temp_img_path)
                
                # 清理臨時文件
                os.remove(temp_img_path)
                os.rmdir(temp_dir)
                
                if text and not text.startswith("提取文本時出錯"):
                    paragraphs.append({
                        'index': page_num,
                        'content': text.strip(),
                        'page': page_num + 1
                    })
        
        doc.close()
        return paragraphs
    
    def display_content(self, word_content, pdf_content):
        """顯示內容預覽"""
        st.title("內容預覽")
        
        # 建立側邊欄
        with st.sidebar:
            st.title("內容統計資訊")
            st.write(f"Word 段落數: {len(word_content)}")
            st.write(f"PDF 頁數: {len(pdf_content)}")
            
            # 如果沒有內容，提供重新提取的選項
            if not word_content or not pdf_content:
                return st.button("重新提取內容", key="refresh_content_btn")
        
        # 建立兩個分頁
        tab1, tab2 = st.tabs(["Word 內容", "PDF 內容"])
        
        with tab1:
            st.subheader("Word 文件內容")
            if not word_content:
                st.warning("Word 文件中沒有找到任何內容")
            else:
                for para in word_content:
                    # 根據段落類型使用不同的顯示方式
                    if para.get('type') == 'heading':
                        st.markdown(f"### {para['content']}")
                    elif para.get('type') == 'toc':
                        st.markdown(f"**目錄項**: {para['content']}")
                    elif para.get('type') == 'table_text':
                        st.markdown(f"**表格內容**:")
                        st.text_area("", para['content'], height=100, key=f"word_para_{para['index']}")
                    else:
                        st.write(f"段落 {para['index'] + 1}:")
                        st.text_area("", para['content'], height=100, key=f"word_para_{para['index']}")
        
        with tab2:
            st.subheader("PDF 文件內容")
            if not pdf_content:
                st.warning("PDF 文件中沒有找到任何內容")
            else:
                for para in pdf_content:
                    st.write(f"頁碼 {para.get('page', para['index'] + 1)}:")
                    st.text_area("", para['content'], height=100, key=f"pdf_para_{para['index']}")
        
        return False 